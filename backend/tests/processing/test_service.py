"""Document processing service tests."""

from __future__ import annotations

from pathlib import Path

import fitz
import pytest
from docx import Document

from rag_enterprise.processing.exceptions import EmptyContentError
from rag_enterprise.processing.service import DocumentProcessingService


@pytest.fixture
def service() -> DocumentProcessingService:
    return DocumentProcessingService()


@pytest.fixture
def persian_txt(tmp_path: Path) -> Path:
    path = tmp_path / "persian.txt"
    path.write_text(
        (
            "علي كتاب یک نمونه متن فارسی است که برای آزمایش "
            "استخراج و نرمال‌سازی نوشته شده است.\n\nمتن دوم"
        ),
        encoding="utf-8",
    )
    return path


@pytest.fixture
def mixed_txt(tmp_path: Path) -> Path:
    path = tmp_path / "mixed.txt"
    path.write_text(
        "این متن فارسی است و this part is English for mixed language detection testing.",
        encoding="utf-8",
    )
    return path


@pytest.fixture
def docx_file(tmp_path: Path) -> Path:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("علي كتاب")
    document.add_paragraph("Paragraph two")
    document.save(path)
    return path


@pytest.fixture
def pdf_file(tmp_path: Path) -> Path:
    path = tmp_path / "sample.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "علي كتاب\n\nPDF paragraph")
    document.save(path)
    document.close()
    return path


def test_processes_persian_text(service: DocumentProcessingService, persian_txt: Path) -> None:
    result = service.process_file(persian_txt)

    assert "علی" in result.text
    assert "ک" in result.text
    assert result.metadata.language == "fa"
    assert result.metadata.parser == "text"
    assert result.metadata.character_count > 0


def test_processes_mixed_persian_english(
    service: DocumentProcessingService,
    mixed_txt: Path,
) -> None:
    result = service.process_file(mixed_txt)

    assert "فارسی" in result.text
    assert "English" in result.text
    assert result.metadata.language in {"fa", "en", "unknown"}


def test_processes_docx(service: DocumentProcessingService, docx_file: Path) -> None:
    result = service.process_file(docx_file)

    assert "علی" in result.text or "كتاب" in result.text
    assert result.metadata.parser == "docx"
    assert "Paragraph two" in result.text


def test_processes_pdf(service: DocumentProcessingService, pdf_file: Path) -> None:
    result = service.process_file(pdf_file)

    assert "PDF paragraph" in result.text
    assert result.metadata.parser == "pdf"
    assert result.metadata.page_count == 1


def test_empty_file_raises(service: DocumentProcessingService, tmp_path: Path) -> None:
    path = tmp_path / "empty.txt"
    path.write_text("   \n\n  ", encoding="utf-8")

    with pytest.raises(EmptyContentError):
        service.process_file(path)
