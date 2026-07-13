"""DOCX parser."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from rag_enterprise.processing.exceptions import CorruptFileError
from rag_enterprise.processing.models import ParserOutput


def parse_docx(path: Path) -> ParserOutput:
    """Extract paragraph and table text from DOCX files."""
    try:
        document = Document(str(path))
    except PackageNotFoundError as exc:
        raise CorruptFileError(f"Cannot parse DOCX file: {path}") from exc
    except OSError as exc:
        raise CorruptFileError(f"Cannot read DOCX file: {path}") from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        line = paragraph.text.strip()
        if line:
            parts.append(line)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return ParserOutput(text="\n\n".join(parts), parser="docx")
